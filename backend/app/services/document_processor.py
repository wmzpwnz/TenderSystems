"""
Обработчик документов тендеров (PDF, DOC, Excel)
"""
import io
from pathlib import Path
from typing import Dict, List, Optional
import PyPDF2
import pdfplumber
from docx import Document
import openpyxl
from PIL import Image
import pytesseract
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Обработчик различных типов документов"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.doc': self._extract_docx,
            '.docx': self._extract_docx,
            '.xls': self._extract_excel,
            '.xlsx': self._extract_excel,
            '.txt': self._extract_text
        }
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Извлечь текст из документа
        
        Args:
            file_content: Байты файла
            filename: Имя файла с расширением
        
        Returns:
            Извлеченный текст
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.supported_formats:
            logger.warning(f"Unsupported file format: {file_ext}")
            return ""
        
        try:
            extractor = self.supported_formats[file_ext]
            text = await extractor(file_content)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return ""
    
    async def _extract_pdf(self, content: bytes) -> str:
        """Извлечение текста из PDF"""
        text_parts = []
        
        # Метод 1: pdfplumber (лучше для структурированных PDF)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Метод 2: PyPDF2 (резервный)
        if not text_parts:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
        
        # Метод 3: OCR для сканов (если текст не извлечен)
        if not text_parts:
            try:
                from pdf2image import convert_from_bytes
                from PIL import ImageEnhance, ImageFilter
                
                logger.info("Attempting OCR for scanned PDF")
                
                # Конвертируем PDF в изображения (300 DPI для лучшего качества)
                images = convert_from_bytes(content, dpi=300, fmt='png')
                
                ocr_texts = []
                for i, img in enumerate(images):
                    try:
                        # Предобработка изображения для улучшения OCR
                        # Конвертируем в grayscale если нужно
                        if img.mode != 'L':
                            img = img.convert('L')
                        
                        # Улучшаем контраст
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(1.5)
                        
                        # Улучшаем резкость
                        img = img.filter(ImageFilter.SHARPEN)
                        
                        # OCR с русским и английским языками
                        page_text = pytesseract.image_to_string(
                            img,
                            lang='rus+eng',
                            config='--psm 6 -c tessedit_char_whitelist=0123456789АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯабвгдеёжзийклмнопрстуфхцчшщъыьэюяABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,;:!?()[]{}/-+=*&%$#@ '
                        )
                        
                        if page_text.strip():
                            ocr_texts.append(f"Страница {i+1}:\n{page_text}")
                            logger.info(f"OCR extracted {len(page_text)} characters from page {i+1}")
                    
                    except Exception as page_error:
                        logger.warning(f"OCR failed for page {i+1}: {page_error}")
                        continue
                
                if ocr_texts:
                    logger.info(f"OCR successfully extracted text from {len(ocr_texts)} pages")
                    return "\n\n".join(ocr_texts)
                else:
                    logger.warning("OCR did not extract any text")
            
            except ImportError:
                logger.warning("pdf2image not installed, skipping OCR")
            except Exception as e:
                logger.warning(f"OCR failed: {e}")
                import traceback
                logger.debug(traceback.format_exc())
        
        return "\n\n".join(text_parts)
    
    async def _extract_docx(self, content: bytes) -> str:
        """Извлечение текста из DOCX"""
        try:
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""
    
    async def _extract_excel(self, content: bytes) -> str:
        """Извлечение текста из Excel"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                if sheet_text:
                    text_parts.append(f"Лист: {sheet_name}\n" + "\n".join(sheet_text))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting Excel: {e}")
            return ""
    
    async def _extract_text(self, content: bytes) -> str:
        """Извлечение текста из TXT"""
        try:
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return ""
    
    async def process_multiple_documents(
        self,
        documents: List[Dict[str, bytes]]
    ) -> str:
        """
        Обработать несколько документов и объединить текст
        
        Args:
            documents: Список словарей с ключами 'filename' и 'content'
        
        Returns:
            Объединенный текст всех документов
        """
        all_texts = []
        
        for doc in documents:
            filename = doc.get('filename', 'unknown')
            content = doc.get('content', b'')
            
            text = await self.extract_text(content, filename)
            if text.strip():
                all_texts.append(f"=== {filename} ===\n{text}")
        
        return "\n\n".join(all_texts)


